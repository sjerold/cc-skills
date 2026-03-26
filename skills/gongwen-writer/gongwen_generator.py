# -*- coding: utf-8 -*-
"""
公文Word文档生成器
根据党政机关公文格式规范生成Word文档
"""

from docx import Document
from docx.shared import Pt, Mm, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import datetime
import os

# 默认输出目录：用户下载目录
DEFAULT_OUTPUT_DIR = r'C:\Users\admin\Downloads'


def get_output_path(filename):
    """获取完整输出路径，默认保存到下载目录"""
    if os.path.isabs(filename):
        return filename
    return os.path.join(DEFAULT_OUTPUT_DIR, filename)

class GongwenDocument:
    """公文文档生成器"""

    # 字号与磅值对应关系
    FONT_SIZES = {
        '二号': Pt(22),      # 方正小标宋
        '三号': Pt(16),      # 仿宋、黑体、楷体
        '四号': Pt(14),      # 版记
    }

    # 首行缩进：2个中文字符
    # 版心156mm/28字 ≈ 5.57mm/字，2字 ≈ 11mm
    INDENT_2_CHARS = Mm(11)

    def __init__(self):
        """初始化文档"""
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """设置页面格式"""
        for section in self.doc.sections:
            # A4纸张
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            # 页边距
            section.top_margin = Mm(36)
            section.bottom_margin = Mm(36)
            section.left_margin = Mm(27)
            section.right_margin = Mm(27)

    def _set_font(self, run, font_name, font_size, bold=False):
        """设置字体"""
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _add_paragraph(self, text, font_name='仿宋_GB2312', font_size=None,
                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                       space_before=Pt(0), space_after=Pt(0),
                       line_spacing=Pt(28), first_line_indent=None,
                       bold=False):
        """添加段落"""
        if font_size is None:
            font_size = self.FONT_SIZES['三号']

        para = self.doc.add_paragraph()
        para.alignment = alignment
        para.paragraph_format.space_before = space_before
        para.paragraph_format.space_after = space_after
        para.paragraph_format.line_spacing = line_spacing
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        if first_line_indent:
            para.paragraph_format.first_line_indent = first_line_indent

        run = para.add_run(text)
        self._set_font(run, font_name, font_size, bold)

        return para

    def add_fenhao(self, fenhao):
        """添加份号"""
        self._add_paragraph(
            fenhao,
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['三号'],
            alignment=WD_ALIGN_PARAGRAPH.LEFT
        )

    def add_miji(self, miji, baomi_qixian=None):
        """添加密级和保密期限"""
        if baomi_qixian:
            text = f"{miji}★{baomi_qixian}"
        else:
            # 只标密级，两字间空1字
            text = f"{miji[0]}　{miji[1]}" if len(miji) == 2 else miji

        self._add_paragraph(
            text,
            font_name='黑体',
            font_size=self.FONT_SIZES['三号'],
            alignment=WD_ALIGN_PARAGRAPH.LEFT
        )

    def add_bantou(self, org_name, has_wenjian=True):
        """添加发文机关标志(版头)"""
        text = f"{org_name}文件" if has_wenjian else org_name

        para = self._add_paragraph(
            text,
            font_name='方正小标宋简体',
            font_size=Pt(26),  # 小于22mm×15mm
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(50),  # 上边缘至上页边距50mm
            space_after=Pt(16)
        )
        return para

    def add_fawenzihao(self, dai_zi, year, order_num):
        """添加发文字号"""
        text = f"{dai_zi}〔{year}〕{order_num}号"

        # 先添加空行
        self.doc.add_paragraph()

        self._add_paragraph(
            text,
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['三号'],
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

    def add_red_line(self, has_star=True):
        """添加红色分隔线"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 创建红色分隔线
        if has_star:
            # 五角星加横线
            run = para.add_run("★")
            run.font.color.rgb = None  # 需要设置为红色

    def add_title(self, org_name, subject, doc_type):
        """添加公文标题"""
        title = f"{org_name}{subject}{doc_type}"

        # 红色分隔线下空一行
        self.doc.add_paragraph()

        self._add_paragraph(
            title,
            font_name='方正小标宋简体',
            font_size=self.FONT_SIZES['二号'],
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(16)
        )

    def add_zhusong_jiguan(self, organs):
        """添加主送机关"""
        # 标题下空一行
        self.doc.add_paragraph()

        if isinstance(organs, list):
            text = '、'.join(organs)
        else:
            text = organs

        text += '：'

        self._add_paragraph(
            text,
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['三号'],
            alignment=WD_ALIGN_PARAGRAPH.LEFT
        )

    def add_body(self, content):
        """添加正文"""
        # 处理正文内容
        lines = content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 判断标题级别
            if line.startswith('一、') or line.startswith('二、') or \
               line.startswith('三、') or line.startswith('四、') or \
               line.startswith('五、') or line.startswith('六、') or \
               line.startswith('七、') or line.startswith('八、') or \
               line.startswith('九、') or line.startswith('十、'):
                # 一级标题：3号黑体
                self._add_paragraph(
                    line,
                    font_name='黑体',
                    font_size=self.FONT_SIZES['三号'],
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent=self.INDENT_2_CHARS  # 左空二字
                )
            elif line.startswith('（一）') or line.startswith('（二）') or \
                 line.startswith('（三）') or line.startswith('（四）') or \
                 line.startswith('（五）') or line.startswith('（六）') or \
                 line.startswith('（七）') or line.startswith('（八）') or \
                 line.startswith('（九）') or line.startswith('（十）'):
                # 二级标题：3号楷体不加粗
                self._add_paragraph(
                    line,
                    font_name='楷体_GB2312',
                    font_size=self.FONT_SIZES['三号'],
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent=self.INDENT_2_CHARS
                )
            elif line[0].isdigit() and '. ' in line[:3]:
                # 三级标题：3号仿宋加粗
                self._add_paragraph(
                    line,
                    font_name='仿宋_GB2312',
                    font_size=self.FONT_SIZES['三号'],
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent=self.INDENT_2_CHARS,
                    bold=True
                )
            else:
                # 普通正文：3号仿宋
                self._add_paragraph(
                    line,
                    font_name='仿宋_GB2312',
                    font_size=self.FONT_SIZES['三号'],
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    first_line_indent=self.INDENT_2_CHARS
                )

    def add_fujian(self, fujian_list):
        """添加附件"""
        # 正文下空一行
        self.doc.add_paragraph()

        # 附件标题
        self._add_paragraph(
            '附件',
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['三号'],
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=self.INDENT_2_CHARS
        )

        # 附件列表
        for i, item in enumerate(fujian_list, 1):
            text = f"{i}. {item}" if len(fujian_list) > 1 else item
            para = self._add_paragraph(
                text,
                font_name='仿宋_GB2312',
                font_size=self.FONT_SIZES['三号'],
                alignment=WD_ALIGN_PARAGRAPH.LEFT
            )
            # 左空二字后对齐（附件：后对齐，约3字位置）
            para.paragraph_format.left_indent = Mm(16)  # 约3字

    def add_fuzhu(self, fuzhu_text):
        """添加附注"""
        self._add_paragraph(
            f"（{fuzhu_text}）",
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['三号'],
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_line_indent=self.INDENT_2_CHARS
        )

    def add_shuming(self, org_name):
        """添加发文机关署名"""
        # 上距正文二至三行
        self.doc.add_paragraph()
        self.doc.add_paragraph()

        self._add_paragraph(
            org_name,
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['三号'],
            alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )

    def add_chengwen_riqi(self, date=None, you_kong=4):
        """添加成文日期"""
        if date is None:
            date = datetime.date.today()

        text = f"{date.year}年{date.month}月{date.day}日"

        para = self._add_paragraph(
            text,
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['三号'],
            alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )

        # 右空4字编排
        para.paragraph_format.right_indent = Mm(21 * 4 / 28 * 10)  # 近似4字

    def add_banji(self, yinfa_jiguan, yinfa_riqi, chaosong=None):
        """添加版记"""
        # 添加分隔线
        para = self.doc.add_paragraph()
        para.add_run('_' * 50)

        # 印发机关和印发日期
        p1 = self._add_paragraph(
            f" {yinfa_jiguan}",
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['四号'],
            alignment=WD_ALIGN_PARAGRAPH.LEFT
        )

        # 印发日期(右对齐)
        if yinfa_riqi is None:
            yinfa_riqi = datetime.date.today()
        date_text = f"{yinfa_riqi.year}年{yinfa_riqi.month}月{yinfa_riqi.day}日印发"

        self._add_paragraph(
            f"                    {date_text} ",
            font_name='仿宋_GB2312',
            font_size=self.FONT_SIZES['四号'],
            alignment=WD_ALIGN_PARAGRAPH.RIGHT
        )

        # 抄送机关
        if chaosong:
            self._add_paragraph(
                f" 抄送：{chaosong}。",
                font_name='仿宋_GB2312',
                font_size=self.FONT_SIZES['四号'],
                alignment=WD_ALIGN_PARAGRAPH.LEFT
            )

    def add_page_number(self):
        """添加页码（4号宋体，阿拉伯数字Times New Roman，一字线）"""
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            # 清除默认段落
            for para in footer.paragraphs:
                para.clear()

            # 使用第一个段落
            para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加一字线
            run1 = para.add_run('— ')
            run1.font.name = '宋体'
            run1.font.size = self.FONT_SIZES['四号']
            run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 添加页码域（阿拉伯数字用Times New Roman）
            fld_char_begin = OxmlElement('w:fldChar')
            fld_char_begin.set(qn('w:fldCharType'), 'begin')

            instr_text = OxmlElement('w:instrText')
            instr_text.text = 'PAGE'

            fld_char_sep = OxmlElement('w:fldChar')
            fld_char_sep.set(qn('w:fldCharType'), 'separate')

            fld_char_end = OxmlElement('w:fldChar')
            fld_char_end.set(qn('w:fldCharType'), 'end')

            # 页码run（Times New Roman字体）
            run2 = para.add_run()
            run2.font.name = 'Times New Roman'
            run2.font.size = self.FONT_SIZES['四号']
            run2._r.append(fld_char_begin)
            run2._r.append(instr_text)
            run2._r.append(fld_char_sep)
            run2._r.append(fld_char_end)

            # 添加后一字线
            run3 = para.add_run(' —')
            run3.font.name = '宋体'
            run3.font.size = self.FONT_SIZES['四号']
            run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 设置页脚距版心距离（约7mm）
            section.footer_distance = Mm(7)

    def save(self, file_path):
        """保存文档"""
        # 在保存前添加页码
        self.add_page_number()
        self.doc.save(file_path)


def create_gongwen(
    output_path: str,
    org_name: str,
    doc_type: str,
    subject: str,
    content: str,
    fawen_zihao: dict = None,
    zhusong: list = None,
    fujian: list = None,
    fuzhu: str = None,
    chaosong: str = None,
    miji: str = None,
    baomi_qixian: str = None,
    fenhao: str = None,
    date: datetime.date = None,
    yinfa_jiguan: str = None
):
    """
    创建公文Word文档

    参数:
        output_path: 输出文件路径
        org_name: 发文机关名称
        doc_type: 文种(通知、报告、请示等)
        subject: 公文主题
        content: 正文内容
        fawen_zihao: 发文字号 {'dai_zi': '苏银发', 'year': 2026, 'order_num': 15}
        zhusong: 主送机关列表
        fujian: 附件列表
        fuzhu: 附注内容
        chaosong: 抄送机关
        miji: 密级
        baomi_qixian: 保密期限
        fenhao: 份号
        date: 成文日期
        yinfa_jiguan: 印发机关
    """
    doc = GongwenDocument()

    # 份号
    if fenhao:
        doc.add_fenhao(fenhao)

    # 密级
    if miji:
        doc.add_miji(miji, baomi_qixian)

    # 版头
    doc.add_bantou(org_name)

    # 发文字号
    if fawen_zihao:
        doc.add_fawenzihao(
            fawen_zihao.get('dai_zi', ''),
            fawen_zihao.get('year', datetime.date.today().year),
            fawen_zihao.get('order_num', 1)
        )

    # 标题
    doc.add_title(org_name, subject, doc_type)

    # 主送机关
    if zhusong:
        doc.add_zhusong_jiguan(zhusong)

    # 正文
    doc.add_body(content)

    # 附件
    if fujian:
        doc.add_fujian(fujian)

    # 附注
    if fuzhu:
        doc.add_fuzhu(fuzhu)

    # 署名
    doc.add_shuming(org_name)

    # 成文日期
    doc.add_chengwen_riqi(date)

    # 版记
    if yinfa_jiguan is None:
        yinfa_jiguan = f"{org_name}办公室"
    doc.add_banji(yinfa_jiguan, date, chaosong)

    # 保存
    output_path = get_output_path(output_path)
    doc.save(output_path)
    print(f"公文已生成: {output_path}")


# 使用示例
if __name__ == '__main__':
    create_gongwen(
        output_path='示例公文.docx',  # 将保存到 C:\Users\admin\Downloads\示例公文.docx
        org_name='苏州银行',
        doc_type='通知',
        subject='关于开展2026年度员工培训的',
        content='''
        各部门、各分行：

        为进一步提升员工专业素质，现将2026年度员工培训工作有关事项通知如下：

        一、培训目标

        通过系统培训，提升员工业务能力和综合素质，促进银行高质量发展。

        二、培训内容

        （一）业务知识培训。包括银行业务基础知识、金融产品知识、风险防控知识等。

        （二）技能培训。包括服务技能、沟通技能、操作技能等。

        三、培训时间

        2026年4月至12月，具体安排见附件。

        四、工作要求

        各部门要高度重视，认真组织，确保培训工作顺利开展。

        特此通知。
        ''',
        fawen_zihao={'dai_zi': '苏银发', 'year': 2026, 'order_num': 25},
        zhusong=['各部门', '各分行'],
        fujian=['2026年度员工培训计划表'],
        fuzhu='此件公开发布',
        chaosong='总行领导班子成员',
        yinfa_jiguan='苏州银行办公室'
    )