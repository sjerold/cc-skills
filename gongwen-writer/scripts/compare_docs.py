# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

print("=" * 60)
print("【源文件内容摘要】")
print("=" * 60)

doc_source = Document(r'C:\Users\admin\Downloads\江苏金服与苏州银行共建数字人民币运营机构项目可行性研究报告.docx')

print(f"段落总数: {len(doc_source.paragraphs)}")
print(f"表格总数: {len(doc_source.tables)}")

# 统计主要内容
content_count = 0
for para in doc_source.paragraphs:
    if para.text.strip() and len(para.text.strip()) > 20:
        content_count += 1

print(f"有效内容段落: {content_count}")

print("\n前20个有效段落:")
count = 0
for para in doc_source.paragraphs:
    if para.text.strip():
        count += 1
        if count <= 20:
            print(f"  [{count}] {para.text.strip()[:60]}...")

print("\n" + "=" * 60)
print("【新公文格式文件内容摘要】")
print("=" * 60)

doc_new = Document(r'C:\Users\admin\Downloads\江苏金服与苏州银行共建数字人民币运营机构项目可行性研究报告（公文格式）.docx')

print(f"段落总数: {len(doc_new.paragraphs)}")
print(f"表格总数: {len(doc_new.tables)}")

print("\n全部段落内容:")
for i, para in enumerate(doc_new.paragraphs):
    if para.text.strip():
        print(f"  [{i}] {para.text.strip()[:80]}...")

print("\n" + "=" * 60)
print("【内容对比】")
print("=" * 60)

# 检查关键内容是否保留
source_text = "\n".join([p.text for p in doc_source.paragraphs])
new_text = "\n".join([p.text for p in doc_new.paragraphs])

keywords = [
    "数字人民币",
    "苏州银行",
    "江苏金服",
    "运营机构",
    "2.0层",
    "钱包",
    "可行性研究"
]

print("关键词保留检查:")
for kw in keywords:
    in_source = kw in source_text
    in_new = kw in new_text
    status = "保留" if (in_source and in_new) else ("丢失" if in_source and not in_new else "无")
    print(f"  '{kw}': {status}")