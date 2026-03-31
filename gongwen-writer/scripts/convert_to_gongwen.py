# -*- coding: utf-8 -*-
"""
将可行性研究报告转换为公文格式 - 完整保留原内容
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Mm, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(run, font_name, size_pt):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_format(para, line_spacing_pt=28, space_before=0, space_after=0):
    """设置段落格式"""
    para.paragraph_format.line_spacing = Pt(line_spacing_pt)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)

def add_heading_paragraph(doc, text, level=1):
    """添加标题段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)

    if level == 1:
        set_chinese_font(run, '黑体', 16)
        run.bold = True
    elif level == 2:
        set_chinese_font(run, '楷体', 16)
    elif level == 3:
        set_chinese_font(run, '仿宋', 16)
        run.bold = True
    else:
        set_chinese_font(run, '仿宋', 16)

    set_paragraph_format(para, line_spacing_pt=28)
    para.paragraph_format.first_line_indent = Pt(0) if level <= 2 else Pt(32)
    return para

def add_body_paragraph(doc, text, indent=True):
    """添加正文段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '仿宋', 16)
    set_paragraph_format(para, line_spacing_pt=28)
    if indent:
        para.paragraph_format.first_line_indent = Pt(32)
    return para

def add_table_content_as_text(doc, table, table_title=None):
    """将表格内容转换为正文段落"""
    if table_title:
        add_body_paragraph(doc, table_title, indent=False)

    for row in table.rows:
        row_text = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                row_text.append(cell_text)
        if row_text:
            line = ' | '.join(row_text)
            add_body_paragraph(doc, line, indent=True)

def create_gongwen_document():
    """创建公文格式文档 - 完整保留原内容"""
    # 读取源文件
    source_path = r'C:\Users\admin\Downloads\江苏金服与苏州银行共建数字人民币运营机构项目可行性研究报告.docx'
    doc_source = Document(source_path)

    # 创建新文档
    doc = Document()

    # 设置页面格式
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Mm(36)
        section.bottom_margin = Mm(36)
        section.left_margin = Mm(27)
        section.right_margin = Mm(27)

    # ===== 版头部分 =====
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("江苏金服数字有限公司文件")
    run.bold = True
    set_chinese_font(run, '方正小标宋简体', 22)

    for _ in range(2):
        doc.add_paragraph()

    fw_para = doc.add_paragraph()
    fw_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fw_para.add_run("苏金发〔2026〕1号")
    set_chinese_font(run, '仿宋', 16)

    doc.add_paragraph()

    # ===== 标题 =====
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("关于江苏金服与苏州银行共建数字人民币运营机构项目可行性研究报告")
    run.bold = True
    set_chinese_font(run, '方正小标宋简体', 22)

    doc.add_paragraph()

    # ===== 主送机关 =====
    main_para = doc.add_paragraph()
    run = main_para.add_run("苏州市人民政府金融工作办公室：")
    set_chinese_font(run, '仿宋', 16)
    set_paragraph_format(main_para, line_spacing_pt=28)

    # ===== 正文 - 完整复制源文件内容 =====
    # 跳过源文件的前几个标题段落（已经在公文格式中处理）
    skip_until = False

    for i, para in enumerate(doc_source.paragraphs):
        text = para.text.strip()

        # 跳过空段落
        if not text:
            continue

        # 跳过源文件的封面信息
        if i < 24:
            # 检测是否到达正文开始
            if text == "项目必要性分析":
                skip_until = True
                add_heading_paragraph(doc, "一、项目必要性分析", level=1)
                continue
            continue

        # 处理正文内容
        if text == "项目必要性分析":
            add_heading_paragraph(doc, "一、项目必要性分析", level=1)
        elif text == "项目实施背景和依据":
            add_heading_paragraph(doc, "（一）项目实施背景和依据", level=2)
        elif text == "实施单位信息化现状":
            add_heading_paragraph(doc, "（二）实施单位信息化现状", level=2)
        elif text == "人力资源与组织架构方面":
            add_heading_paragraph(doc, "1.人力资源与组织架构方面", level=3)
        elif text == "研发能力与技术能力方面":
            add_heading_paragraph(doc, "2.研发能力与技术能力方面", level=3)
        elif text == "知识产权与平台产品方面":
            add_heading_paragraph(doc, "3.知识产权与平台产品方面", level=3)
        elif text == "存在主要问题与差距":
            add_heading_paragraph(doc, "4.存在主要问题与差距", level=3)
        elif text == "项目实施必要性分析":
            add_heading_paragraph(doc, "（三）项目实施必要性分析", level=2)
        elif text == "往期项目的成果成效":
            add_heading_paragraph(doc, "（四）往期项目的成果成效", level=2)
        elif text == "无":
            add_body_paragraph(doc, "无。")
        elif text == "项目需求分析":
            add_heading_paragraph(doc, "二、项目需求分析", level=1)
        elif text == "业务分析":
            add_heading_paragraph(doc, "（一）业务分析", level=2)
        elif text == "目标分析":
            add_heading_paragraph(doc, "（二）目标分析", level=2)
        elif text == "功能和性能分析":
            add_heading_paragraph(doc, "（三）功能和性能分析", level=2)
        elif text == "钱包服务功能":
            add_heading_paragraph(doc, "（四）钱包服务功能", level=2)
        elif text == "业务产品功能":
            add_heading_paragraph(doc, "（五）业务产品功能", level=2)
        elif text == "零售业务相关功能":
            add_heading_paragraph(doc, "1.零售业务相关功能", level=3)
        elif text == "公司业务相关功能":
            add_heading_paragraph(doc, "2.公司业务相关功能", level=3)
        elif text == "信贷业务相关功能":
            add_heading_paragraph(doc, "3.信贷业务相关功能", level=3)
        elif text == "收单服务功能":
            add_heading_paragraph(doc, "（六）收单服务功能", level=2)
        elif text == "场景接入功能":
            add_heading_paragraph(doc, "（七）场景接入功能", level=2)
        elif text == "风控管理功能":
            add_heading_paragraph(doc, "（八）风控管理功能", level=2)
        elif text == "服务输出功能":
            add_heading_paragraph(doc, "（九）服务输出功能", level=2)
        elif text == "数币额度流动性管理功能":
            add_heading_paragraph(doc, "（十）数币额度流动性管理功能", level=2)
        elif text == "数币结算钱包额度管理":
            add_heading_paragraph(doc, "（十一）数币结算钱包额度管理", level=2)
        elif text == "软件建设方案":
            add_heading_paragraph(doc, "三、软件建设方案", level=1)
        elif text == "系统架构":
            add_heading_paragraph(doc, "（一）系统架构", level=2)
        elif text == "总体架构":
            add_heading_paragraph(doc, "1.总体架构", level=3)
        elif text == "应用架构":
            add_heading_paragraph(doc, "2.应用架构", level=3)
        elif text == "网络架构":
            add_heading_paragraph(doc, "3.网络架构", level=3)
        elif text == "入口说明":
            add_heading_paragraph(doc, "（二）入口说明", level=2)
        elif text == "不涉及":
            add_body_paragraph(doc, "不涉及。")
        elif text == "功能说明":
            add_heading_paragraph(doc, "（三）功能说明", level=2)
        elif text == "用户体系":
            add_heading_paragraph(doc, "1.用户体系", level=3)
        elif text == "整合对接":
            add_heading_paragraph(doc, "（四）整合对接", level=2)
        elif text == "市级政务云需求说明":
            add_heading_paragraph(doc, "1.市级政务云需求说明", level=3)
        elif text == "公共能力复用需求说明":
            add_heading_paragraph(doc, "2.公共能力复用需求说明", level=3)
        elif text == "内外部系统平台对接说明":
            add_heading_paragraph(doc, "3.内外部系统平台对接说明", level=3)
        elif text == "政务信息资源目录预编":
            add_heading_paragraph(doc, "4.政务信息资源目录预编", level=3)
        elif text == "网络安全保障体系":
            add_heading_paragraph(doc, "（五）网络安全保障体系", level=2)
        elif text == "商用密码应用与安全性评估":
            add_heading_paragraph(doc, "（六）商用密码应用与安全性评估", level=2)
        elif text == "硬件采购方案":
            add_heading_paragraph(doc, "四、硬件采购方案", level=1)
        elif text == "采购类型":
            add_heading_paragraph(doc, "（一）采购类型", level=2)
        elif text == "市场调研分析":
            add_heading_paragraph(doc, "（二）市场调研分析", level=2)
        elif text == "采购清单":
            add_heading_paragraph(doc, "（三）采购清单", level=2)
        elif text == "部署位置":
            add_heading_paragraph(doc, "（四）部署位置", level=2)
        elif text == "服务方案":
            add_heading_paragraph(doc, "五、服务方案", level=1)
        elif text == "服务内容":
            add_heading_paragraph(doc, "（一）服务内容", level=2)
        elif text == "服务产出":
            add_heading_paragraph(doc, "（二）服务产出", level=2)
        elif text == "考核指标":
            add_heading_paragraph(doc, "（三）考核指标", level=2)
        elif text == "运行管理体系":
            add_heading_paragraph(doc, "六、运行管理体系", level=1)
        elif text == "运维方案":
            add_heading_paragraph(doc, "（一）运维方案", level=2)
        elif text == "运营方案":
            add_heading_paragraph(doc, "（二）运营方案", level=2)
        elif text == "安全可靠产品和技术选用":
            add_heading_paragraph(doc, "（三）安全可靠产品和技术选用", level=2)
        elif text == "实施进度安排":
            add_heading_paragraph(doc, "七、实施进度安排", level=1)
        elif text == "投资估算":
            add_heading_paragraph(doc, "八、投资估算", level=1)
        elif text == "项目总投资估算":
            add_heading_paragraph(doc, "（一）项目总投资估算", level=2)
        elif text == "项目投资估算明细":
            add_heading_paragraph(doc, "（二）项目投资估算明细", level=2)
        elif text.startswith("表 ") or text.startswith("表1") or text.startswith("表2") or text.startswith("表3") or text.startswith("表4"):
            # 表格标题，跳过，表格内容单独处理
            continue
        elif text.startswith("附件"):
            # 附件信息跳过
            continue
        else:
            # 普通正文段落，完整保留
            add_body_paragraph(doc, text)

    # ===== 处理表格内容 =====
    add_heading_paragraph(doc, "附：系统功能详细说明", level=1)

    # 表格1：数字人民币运营机构系统功能说明
    if len(doc_source.tables) >= 1:
        add_body_paragraph(doc, "表1 数字人民币运营机构系统功能说明", indent=False)
        table1 = doc_source.tables[0]
        for row_idx, row in enumerate(table1.rows):
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    # 处理单元格内的换行
                    cell_text = cell_text.replace('\n', '；')
                    cells_text.append(cell_text)
            if cells_text:
                line = ''.join(cells_text)
                if row_idx == 0:
                    # 表头行
                    add_body_paragraph(doc, f"【{line}】", indent=False)
                else:
                    add_body_paragraph(doc, line)

    # 表格2：网络安全保障体系
    if len(doc_source.tables) >= 2:
        doc.add_paragraph()
        add_body_paragraph(doc, "表2 网络安全保障体系", indent=False)
        table2 = doc_source.tables[1]
        for row_idx, row in enumerate(table2.rows):
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cell_text = cell_text.replace('\n', '；')
                    cells_text.append(cell_text)
            if cells_text:
                line = ''.join(cells_text)
                if row_idx == 0:
                    add_body_paragraph(doc, f"【{line}】", indent=False)
                else:
                    add_body_paragraph(doc, line)

    # 表格3：商用密码应用与安全性评估
    if len(doc_source.tables) >= 3:
        doc.add_paragraph()
        add_body_paragraph(doc, "表3 商用密码应用与安全性评估", indent=False)
        table3 = doc_source.tables[2]
        for row_idx, row in enumerate(table3.rows):
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cell_text = cell_text.replace('\n', '；')
                    cells_text.append(cell_text)
            if cells_text:
                line = ''.join(cells_text)
                if row_idx == 0:
                    add_body_paragraph(doc, f"【{line}】", indent=False)
                else:
                    add_body_paragraph(doc, line)

    # ===== 发文机关署名 =====
    for _ in range(3):
        doc.add_paragraph()

    sign_para = doc.add_paragraph()
    sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sign_para.add_run("江苏金服数字有限公司")
    set_chinese_font(run, '仿宋', 16)
    set_paragraph_format(sign_para, line_spacing_pt=28)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para.add_run("2026年3月31日")
    set_chinese_font(run, '仿宋', 16)
    set_paragraph_format(date_para, line_spacing_pt=28)

    # ===== 版记部分 =====
    for _ in range(2):
        doc.add_paragraph()

    print_para = doc.add_paragraph()
    run = print_para.add_run("印发机关：江苏金服数字有限公司综合管理部                    印发日期：2026年3月31日")
    set_chinese_font(run, '仿宋', 14)

    output_path = r'C:\Users\admin\Downloads\江苏金服苏州银行数字人民币运营机构可研报告_公文格式.docx'
    doc.save(output_path)
    print(f"公文格式文档已保存至: {output_path}")
    return output_path

if __name__ == "__main__":
    create_gongwen_document()