# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'C:\Users\admin\Downloads\江苏金服与苏州银行共建数字人民币运营机构项目可行性研究报告.docx')

print("=== 文档段落 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {para.text}")

print("\n=== 表格内容 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {t_idx + 1} ---")
    for row in table.rows:
        row_text = [cell.text.strip() for cell in row.cells]
        print(' | '.join(row_text))