# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

print('='*60)
print('【源文件统计】')
print('='*60)
doc_source = Document(r'C:\Users\admin\Downloads\江苏金服与苏州银行共建数字人民币运营机构项目可行性研究报告.docx')
print(f'段落总数: {len(doc_source.paragraphs)}')
print(f'表格总数: {len(doc_source.tables)}')

# 统计表格行数
for i, t in enumerate(doc_source.tables):
    print(f'  表格{i+1}: {len(t.rows)}行')

print()
print('='*60)
print('【新公文格式文件统计】')
print('='*60)
doc_new = Document(r'C:\Users\admin\Downloads\江苏金服苏州银行数字人民币运营机构可研报告_公文格式.docx')
print(f'段落总数: {len(doc_new.paragraphs)}')
print(f'表格总数: {len(doc_new.tables)}')

# 统计有效段落
source_chars = sum(len(p.text) for p in doc_source.paragraphs)
new_chars = sum(len(p.text) for p in doc_new.paragraphs)
print(f'源文件字符总数: {source_chars}')
print(f'新文件字符总数: {new_chars}')
print(f'内容保留率: {new_chars/source_chars*100:.1f}%')

# 统计表格字符
table_chars = 0
for t in doc_source.tables:
    for row in t.rows:
        for cell in row.cells:
            table_chars += len(cell.text)
print(f'源文件表格字符数: {table_chars}')

# 检查关键词
source_text = ''.join([p.text for p in doc_source.paragraphs])
new_text = ''.join([p.text for p in doc_new.paragraphs])

print()
print('【关键内容检查】')
keywords = ['数字人民币', '苏州银行', '江苏金服', '运营机构', '钱包管理', '智能合约', '风控管理', '网络架构', '网络安全保障', '商用密码', '投资估算', '实施进度']
for kw in keywords:
    s = kw in source_text
    n = kw in new_text
    status = '保留' if (s and n) else ('丢失' if s and not n else '新增')
    print(f'  {kw}: {status}')