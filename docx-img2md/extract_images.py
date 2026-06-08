"""
从docx中提取所有嵌入图片到文字版/<文档名>/pic目录

用法:
    python extract_images.py <docx文件路径>
"""

import os
import sys
import io
from pathlib import Path

from docx import Document
from PIL import Image


def extract_images(docx_path):
    """从docx中提取所有嵌入图片"""
    docx_path = os.path.abspath(docx_path)
    if not os.path.exists(docx_path):
        print(f"错误: 文件不存在: {docx_path}")
        sys.exit(1)

    doc = Document(docx_path)
    doc_dir = os.path.dirname(docx_path)
    basename = Path(docx_path).stem

    # 每个docx输出到独立的子目录，避免多个docx互相覆盖
    # 输出结构: 文字版/<文档名>/pic/ + 文字版/<文档名>/<文档名>.md
    md_dir = os.path.join(doc_dir, "文字版", basename)
    pic_dir = os.path.join(md_dir, "pic")
    os.makedirs(pic_dir, exist_ok=True)

    seen_embeds = set()
    count = 0

    def save_image(img_part):
        nonlocal count
        count += 1
        img_name = f"image_{count:03d}.png"
        img_path = os.path.join(pic_dir, img_name)
        try:
            img = Image.open(io.BytesIO(img_part.blob)).convert('RGB')
            img.save(img_path, 'PNG')
        except Exception:
            with open(img_path, 'wb') as f:
                f.write(img_part.blob)
        print(f"  [{count}] {img_name}")

    def process_drawing(drawing):
        blips = drawing.findall(
            './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
        )
        for blip in blips:
            embed = blip.get(
                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
            )
            if embed and embed not in seen_embeds:
                seen_embeds.add(embed)
                img_part = doc.part.related_parts.get(embed)
                if img_part:
                    save_image(img_part)

    for para in doc.paragraphs:
        for run in para.runs:
            for drawing in run._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
            ):
                process_drawing(drawing)

    # 也检查表格中的图片
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for drawing in run._element.findall(
                            './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
                        ):
                            process_drawing(drawing)

    print(f"\n共提取 {count} 张图片到: {pic_dir}")
    print(f"md文件应放在: {md_dir}/")
    return count


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python extract_images.py <docx文件路径>")
        sys.exit(1)
    extract_images(sys.argv[1])
