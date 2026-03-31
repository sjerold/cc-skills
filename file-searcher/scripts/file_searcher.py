#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件内容搜索工具 - 优化版
支持在 Word (.docx)、PDF、Markdown、文本等文件中搜索关键词
优化: 多进程并行、文件大小限制、快速PDF提取
"""

import sys
import os
import argparse
import json
import re
import warnings
from pathlib import Path
from concurrent.futures import ProcessPoolExecutor, as_completed
import time

# 设置 UTF-8 输出
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# 忽略 PDF 警告
warnings.filterwarnings('ignore', category=UserWarning, module='PyPDF2')

# 默认搜索路径
DEFAULT_PATH = r"C:\Users\admin\Downloads"

# 性能配置
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB 文件大小限制
PDF_TIMEOUT = 10  # PDF处理超时秒数
MAX_PDF_PAGES = 100  # PDF最大页数限制

# 支持的文件扩展名
SUPPORTED_EXTENSIONS = {
    'text': ['.txt', '.md', '.json', '.csv', '.xml', '.log', '.yaml', '.yml', '.ini', '.cfg'],
    'code': ['.py', '.js', '.ts', '.java', '.go', '.rs', '.c', '.cpp', '.h', '.hpp', '.cs', '.php', '.rb', '.swift', '.kt'],
    'docx': ['.docx'],
    'pdf': ['.pdf'],
}


def extract_text_from_docx(file_path):
    """从 Word 文档提取文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return '\n'.join(text), None
    except ImportError:
        return None, "需要安装 python-docx: pip install python-docx"
    except Exception as e:
        return None, f"读取 Word 文件失败: {str(e)}"


def extract_text_from_pdf(file_path, timeout=PDF_TIMEOUT, max_pages=MAX_PDF_PAGES):
    """从 PDF 文档提取文本 - 优化版，使用 PyMuPDF (fitz) 优先"""
    try:
        # 优先使用 PyMuPDF (fitz)，速度最快
        import fitz  # PyMuPDF
        text = []
        doc = fitz.open(file_path)
        page_count = len(doc)
        # 限制页数
        for page_num in range(min(page_count, max_pages)):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text:
                text.append(f"[第{page_num + 1}页]\n{page_text}")
        doc.close()
        return '\n'.join(text), None
    except ImportError:
        pass
    except Exception as e:
        return None, f"读取 PDF 文件失败: {str(e)}"

    # 回退到 PyPDF2
    try:
        import PyPDF2
        text = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            for page_num in range(min(page_count, max_pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text.append(f"[第{page_num + 1}页]\n{page_text}")
        return '\n'.join(text), None
    except ImportError:
        return None, "需要安装 PyMuPDF (pip install pymupdf) 或 PyPDF2"
    except Exception as e:
        return None, f"读取 PDF 文件失败: {str(e)}"


def read_text_file(file_path):
    """读取文本文件"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read(), None
        except UnicodeDecodeError:
            continue
        except Exception as e:
            return None, f"读取文件失败: {str(e)}"
    return None, "无法识别文件编码"


def get_file_extension(file_path):
    """获取文件扩展名（小写）"""
    return Path(file_path).suffix.lower()


def search_in_text(text, keyword, context_chars=100):
    """在文本中搜索关键词，返回匹配上下文"""
    if not text or not keyword:
        return []

    matches = []
    # 使用正则查找所有匹配位置
    pattern = re.compile(re.escape(keyword), re.IGNORECASE)

    for match in pattern.finditer(text):
        start = max(0, match.start() - context_chars)
        end = min(len(text), match.end() + context_chars)

        # 提取上下文
        context = text[start:end]

        # 清理上下文中的多余空白
        context = re.sub(r'\s+', ' ', context).strip()

        # 高亮关键词
        highlighted = pattern.sub(f'【{keyword}】', context)

        matches.append({
            'position': match.start(),
            'context': highlighted
        })

    return matches


def process_single_file(args):
    """处理单个文件（用于并行处理）"""
    file_path, ext, keyword, max_matches_per_file = args

    # 根据文件类型提取文本
    text = None
    error = None

    if ext in SUPPORTED_EXTENSIONS['docx']:
        text, error = extract_text_from_docx(file_path)
    elif ext in SUPPORTED_EXTENSIONS['pdf']:
        text, error = extract_text_from_pdf(file_path)
    else:
        text, error = read_text_file(file_path)

    if error:
        return {'file': str(file_path), 'error': error, 'ext': ext}

    if text:
        matches = search_in_text(text, keyword)
        if matches:
            return {
                'file': str(file_path),
                'filename': Path(file_path).name,
                'ext': ext,
                'matches': matches[:max_matches_per_file],
                'count': len(matches)
            }
    return None


def search_files(keyword, search_path, extensions=None, max_matches_per_file=3, show_progress=True, max_workers=4):
    """搜索文件 - 优化版，支持并行处理"""
    search_path = Path(search_path)
    if not search_path.exists():
        return {'error': f'路径不存在: {search_path}'}

    # 确定要搜索的扩展名
    if extensions:
        ext_list = [ext if ext.startswith('.') else f'.{ext}' for ext in extensions.split(',')]
    else:
        ext_list = None

    # 支持的所有扩展名
    all_supported = SUPPORTED_EXTENSIONS['text'] + SUPPORTED_EXTENSIONS['code'] + SUPPORTED_EXTENSIONS['docx'] + SUPPORTED_EXTENSIONS['pdf']

    # 第一步：收集所有待处理的文件
    files_to_process = []
    total_files = 0
    skipped_large = 0

    for file_path in search_path.rglob('*'):
        if not file_path.is_file():
            continue
        total_files += 1
        ext = get_file_extension(file_path)

        # 检查扩展名是否在支持范围内
        if ext not in all_supported:
            continue

        # 如果指定了扩展名，检查是否匹配
        if ext_list and ext not in ext_list:
            continue

        # 检查文件大小
        try:
            file_size = file_path.stat().st_size
            if file_size > MAX_FILE_SIZE:
                skipped_large += 1
                continue
        except:
            continue

        files_to_process.append((str(file_path), ext, keyword, max_matches_per_file))

    # 第二步：并行处理文件
    results = []
    errors = []

    # 创建进度条 (使用 stderr 避免 stdout 缓冲问题)
    if show_progress:
        try:
            from tqdm import tqdm
            pbar = tqdm(total=len(files_to_process), desc="搜索进度", unit="文件",
                       bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}]",
                       file=sys.stderr)
        except ImportError:
            pbar = None
            print(f"正在搜索 {len(files_to_process)} 个文件...", file=sys.stderr)
    else:
        pbar = None

    # 使用进程池并行处理
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(process_single_file, args): args[0] for args in files_to_process}

        for future in as_completed(futures):
            if pbar:
                pbar.update(1)

            try:
                result = future.result()
                if result:
                    if 'error' in result:
                        errors.append(result)
                    else:
                        results.append(result)
            except Exception as e:
                pass

    if pbar:
        pbar.close()

    return {
        'keyword': keyword,
        'path': str(search_path),
        'results': results,
        'total_files': total_files,
        'scanned_files': len(files_to_process),
        'matched_files': len(results),
        'skipped_large': skipped_large,
        'errors': errors[:5]
    }


def format_output(result):
    """格式化输出结果"""
    if 'error' in result:
        return f"错误: {result['error']}"

    lines = []
    lines.append("=" * 60)
    lines.append(f"搜索关键词: 【{result['keyword']}】")
    lines.append(f"搜索路径: {result['path']}")
    lines.append(f"扫描文件: {result['scanned_files']} / {result['total_files']} (支持格式的文件 / 总文件)")
    if result.get('skipped_large', 0) > 0:
        lines.append(f"跳过大文件: {result['skipped_large']} 个 (>50MB)")
    lines.append(f"匹配文件: {result['matched_files']} 个")
    lines.append("=" * 60)

    if not result['results']:
        lines.append("未找到匹配内容")
        return '\n'.join(lines)

    for i, file_result in enumerate(result['results'], 1):
        lines.append(f"\n[{i}] {file_result['filename']}")
        lines.append(f"    路径: {file_result['file']}")
        lines.append(f"    匹配次数: {file_result['count']}")
        lines.append(f"    匹配内容:")

        for j, match in enumerate(file_result['matches'], 1):
            lines.append(f"      ({j}) ...{match['context']}...")

    if result['errors']:
        lines.append(f"\n处理出错的文件 ({len(result['errors'])} 个):")
        for err in result['errors']:
            lines.append(f"  - {Path(err['file']).name}: {err['error']}")

    return '\n'.join(lines)


def main():
    parser = argparse.ArgumentParser(description='文件内容搜索工具')
    parser.add_argument('keyword', help='搜索关键词')
    parser.add_argument('--path', '-p', default=DEFAULT_PATH, help='搜索路径')
    parser.add_argument('--ext', '-e', default=None, help='文件扩展名，逗号分隔')
    parser.add_argument('--json', '-j', action='store_true', help='输出 JSON 格式')
    parser.add_argument('--max', '-m', type=int, default=3, help='每个文件最多显示的匹配数')
    parser.add_argument('--workers', '-w', type=int, default=4, help='并行进程数')
    parser.add_argument('--no-progress', action='store_true', help='不显示进度条')

    args = parser.parse_args()

    result = search_files(
        args.keyword,
        args.path,
        args.ext,
        args.max,
        show_progress=not args.no_progress,
        max_workers=args.workers
    )

    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(format_output(result))


if __name__ == '__main__':
    main()